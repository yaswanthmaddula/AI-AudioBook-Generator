import pdfplumber
import pytesseract
import docx
import io

# Tesseract path
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# 🔥 CLEANING FUNCTION (IMPROVED)
def clean_text(text):
    lines = text.split("\n")

    cleaned = []
    seen = set()

    for line in lines:
        line = line.strip()

        # remove garbage, duplicates, short junk
        if (
            len(line) > 5
            and line not in seen
            and not line.lower().startswith(("o ", "=", "»", "▪"))
        ):
            cleaned.append(line)
            seen.add(line)

    return "\n".join(cleaned)


def extract_text(file):
    text = ""
    extension = file.name.split(".")[-1].lower()

    file.seek(0)

    # ---------- PDF ----------
    if extension == "pdf":
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:

                # 1️⃣ Extract normal text
                page_text = page.extract_text()

                if page_text and len(page_text.strip()) > 50:
                    text += page_text + "\n"
                else:
                    # 2️⃣ ONLY use OCR if text is missing
                    try:
                        image = page.to_image(resolution=300).original
                        ocr_text = pytesseract.image_to_string(image)
                        text += ocr_text + "\n"
                    except:
                        pass

    # ---------- DOCX ----------
    elif extension == "docx":
        file_bytes = file.read()
        doc = docx.Document(io.BytesIO(file_bytes))

        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text += row_text + "\n"

    # ---------- TXT ----------
    elif extension == "txt":
        text = file.read().decode("utf-8", errors="ignore")

    return clean_text(text)